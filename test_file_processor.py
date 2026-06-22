#!/usr/bin/env python3
"""Test script for file processor."""
import os
import sys
from pathlib import Path

# Add project to path
sys.path.insert(0, str(Path(__file__).parent))

from app.services.utils.RAG_pipeline.processed_data.processed_data import ProcessedData

def create_test_files():
    """Create test files in data directory."""
    data_dir = Path("app/services/utils/RAG_pipeline/data/new_docs")
    data_dir.mkdir(parents=True, exist_ok=True)

    # Create test TXT
    txt_file = data_dir / "test_sample.txt"
    txt_file.write_text("This is a test text file.\nIt contains multiple lines.\nUsed for testing text loading.")
    print(f"[OK] Created {txt_file}")

    # Create test CSV
    csv_file = data_dir / "test_data.csv"
    csv_file.write_text("name,age,city\nAlice,30,New York\nBob,25,Los Angeles\nCharlie,35,Chicago")
    print(f"[OK] Created {csv_file}")

    # Create test DOCX (using python-docx)
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test DOCX document.")
        doc.add_paragraph("It contains multiple paragraphs for testing.")
        docx_file = data_dir / "test_document.docx"
        doc.save(str(docx_file))
        print(f"[OK] Created {docx_file}")
    except ImportError:
        print("⚠ python-docx not installed, skipping DOCX test")

    # Create test XLSX (using openpyxl)
    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws['A1'] = "Product"
        ws['B1'] = "Price"
        ws['A2'] = "Apple"
        ws['B2'] = 1.50
        ws['A3'] = "Banana"
        ws['B3'] = 0.75
        xlsx_file = data_dir / "test_spreadsheet.xlsx"
        wb.save(str(xlsx_file))
        print(f"[OK] Created {xlsx_file}")
    except ImportError:
        print("⚠ openpyxl not installed, skipping XLSX test")

def test_processor():
    """Test the file processor."""
    print("\n" + "="*60)
    print("Testing File Processor")
    print("="*60 + "\n")

    processor = ProcessedData()

    # Test 1: Load all file paths
    print("Test 1: Loading all file paths...")
    file_paths = processor.load_all_file_paths()
    print(f"Found {len(file_paths)} files\n")

    # Test 2: Load all documents
    print("Test 2: Loading all documents...")
    docs = processor.load_all_documents()
    print(f"\nLoaded {len(docs)} documents")

    if docs:
        print("\nSample document info:")
        for i, doc in enumerate(docs[:3]):
            print(f"\n  Doc {i+1}:")
            print(f"    Type: {doc.metadata.get('doc_type')}")
            print(f"    File: {doc.metadata.get('filename')}")
            print(f"    Content preview: {doc.page_content[:80]}...")

    # Test 3: Chunking
    print("\n\nTest 3: Chunking documents...")
    chunks = processor.chunking()
    print(f"Created {len(chunks)} chunks")

    if chunks:
        print("\nSample chunk info:")
        for i, chunk in enumerate(chunks[:2]):
            print(f"\n  Chunk {i+1}:")
            print(f"    Type: {chunk.metadata.get('doc_type')}")
            print(f"    Char count: {chunk.metadata.get('char_count')}")
            print(f"    Preview: {chunk.metadata.get('preview')}")

    print("\n" + "="*60)
    print("[OK] All tests completed successfully")
    print("="*60)

if __name__ == "__main__":
    create_test_files()
    test_processor()
